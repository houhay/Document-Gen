#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Word 文档导出器
- 政企公文格式排版（仿宋/黑体、页边距、行距）
- 封面页（可选）
- 新闻文章格式
- 文件名生成与输出
"""

import os
import re
from typing import Optional

from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn

from intent_engine import WritingIntent
from utils import sanitize_filename, generate_sequential_filename


class WordExporter:
    """Word 文档导出器"""

    def create_document(self, title: str) -> Document:
        """
        创建符合政企公文格式的空白文档
        :param title: 文档标题
        :return: Document 对象
        """
        doc = Document()

        # 页面设置（政企公文标准）
        section = doc.sections[0]
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(3.18)
        section.right_margin = Cm(2.54)

        # 设置默认字体（支持中文）
        style = doc.styles['Normal']
        font = style.font
        font.name = '仿宋_GB2312'
        font.size = Pt(16)  # 3号字
        style.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')

        pf = style.paragraph_format
        pf.line_spacing = 1.5
        pf.first_line_indent = Pt(32)
        pf.space_before = Pt(0)
        pf.space_after = Pt(0)

        # 标题 1 样式（黑体，小2号，居中）
        h1_style = doc.styles['Heading 1']
        h1_font = h1_style.font
        h1_font.name = '黑体'
        h1_font.size = Pt(18)
        h1_font.bold = True
        h1_style.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
        h1_pf = h1_style.paragraph_format
        h1_pf.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        h1_pf.space_before = Pt(12)
        h1_pf.space_after = Pt(12)

        # 标题 2 样式（黑体，3号，左对齐）
        h2_style = doc.styles['Heading 2']
        h2_font = h2_style.font
        h2_font.name = '黑体'
        h2_font.size = Pt(16)
        h2_font.bold = True
        h2_style.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
        h2_pf = h2_style.paragraph_format
        h2_pf.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        h2_pf.space_before = Pt(12)
        h2_pf.space_after = Pt(6)

        # 标题 3 样式（仿宋加粗，3号）
        h3_style = doc.styles['Heading 3']
        h3_font = h3_style.font
        h3_font.name = '仿宋_GB2312'
        h3_font.size = Pt(16)
        h3_font.bold = True
        h3_style.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
        h3_pf = h3_style.paragraph_format
        h3_pf.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        h3_pf.space_before = Pt(6)
        h3_pf.space_after = Pt(6)

        return doc

    def markdown_to_docx(self, doc: Document, markdown_text: str):
        """
        将 Markdown 内容写入 Word 文档
        :param doc: Document 对象
        :param markdown_text: Markdown 格式的文本
        """
        lines = markdown_text.split('\n')

        for line in lines:
            line_stripped = line.strip()

            if not line_stripped:
                # 空行加一个空段落
                doc.add_paragraph()
                continue

            if line_stripped.startswith('### '):
                p = doc.add_heading(line_stripped[4:], level=3)
                self._set_run_font(p, '仿宋_GB2312', Pt(16), bold=True)

            elif line_stripped.startswith('## '):
                p = doc.add_heading(line_stripped[3:], level=2)
                self._set_run_font(p, '黑体', Pt(16), bold=True)

            elif line_stripped.startswith('# '):
                p = doc.add_heading(line_stripped[2:], level=1)
                self._set_run_font(p, '黑体', Pt(18), bold=True)

            elif line_stripped.startswith('- '):
                p = doc.add_paragraph(line_stripped[2:], style='List Bullet')
                self._set_run_font(p, '仿宋_GB2312', Pt(16))

            elif re.match(r'^\d+\.\s', line_stripped):
                text = re.sub(r'^\d+\.\s*', '', line_stripped)
                p = doc.add_paragraph(text, style='List Number')
                self._set_run_font(p, '仿宋_GB2312', Pt(16))

            elif line_stripped == '---':
                p = doc.add_paragraph()
                run = p.add_run('─' * 40)
                run.font.name = '仿宋_GB2312'
                run.font.size = Pt(16)
                run.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')

            else:
                p = doc.add_paragraph(line)
                self._set_run_font(p, '仿宋_GB2312', Pt(16))

    def _set_run_font(self, paragraph, font_name: str, font_size: Pt, bold: bool = False):
        """设置段落中所有 run 的字体"""
        for run in paragraph.runs:
            run.font.name = font_name
            run.font.size = font_size
            run.font.bold = bold
            run.element.rPr.rFonts.set(qn('w:eastAsia'), font_name)

    def add_cover_page(self, doc: Document, title: str, subtitle: str = "", date_str: str = ""):
        """
        添加封面页
        :param doc: Document 对象
        :param title: 文档标题
        :param subtitle: 副标题
        :param date_str: 日期
        """
        # 添加空行
        for _ in range(6):
            doc.add_paragraph()

        # 标题
        p = doc.add_paragraph()
        p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        run = p.add_run(title)
        run.font.name = '黑体'
        run.font.size = Pt(26)
        run.font.bold = True
        run.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')

        doc.add_paragraph()

        # 副标题
        if subtitle:
            p = doc.add_paragraph()
            p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            run = p.add_run(subtitle)
            run.font.name = '仿宋_GB2312'
            run.font.size = Pt(18)
            run.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')

            doc.add_paragraph()
            doc.add_paragraph()

        # 日期
        if date_str:
            doc.add_paragraph()
            p = doc.add_paragraph()
            p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            run = p.add_run(date_str)
            run.font.name = '仿宋_GB2312'
            run.font.size = Pt(16)
            run.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')

        # 分页
        doc.add_page_break()

    def add_news_article(self, doc: Document, title: str, source: str,
                         publish_date: str, content: str, author: str = ""):
        """
        将新闻文章写入 Word 文档
        :param doc: Document 对象
        :param title: 文章标题
        :param source: 来源
        :param publish_date: 发布日期
        :param content: 正文
        :param author: 作者
        """
        # 标题（黑体居中）
        p = doc.add_heading(title, level=1)
        p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        self._set_run_font(p, '黑体', Pt(18), bold=True)

        # 来源信息行
        info_parts = []
        if source:
            info_parts.append(f"来源：{source}")
        if author:
            info_parts.append(f"作者：{author}")
        if publish_date:
            info_parts.append(f"日期：{publish_date}")

        if info_parts:
            p = doc.add_paragraph()
            p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            run = p.add_run("  ".join(info_parts))
            run.font.name = '仿宋_GB2312'
            run.font.size = Pt(14)
            run.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')

        # 正文
        paragraphs = content.split('\n')
        for para in paragraphs:
            para = para.strip()
            if para:
                p = doc.add_paragraph(para)
                self._set_run_font(p, '仿宋_GB2312', Pt(16))
            else:
                doc.add_paragraph()

    def export(self, content: str, intent: WritingIntent,
               output_dir: str, file_name: str = None,
               add_cover: bool = False) -> str:
        """
        导出文档为 .docx 文件
        :param content: Markdown 内容
        :param intent: 写作意图
        :param output_dir: 输出目录
        :param file_name: 文件名（不含扩展名），None 则自动生成
        :param add_cover: 是否添加封面页
        :return: 保存路径
        """
        os.makedirs(output_dir, exist_ok=True)

        # 创建文档
        doc = self.create_document(intent.topic)

        # 封面
        if add_cover:
            import datetime
            self.add_cover_page(
                doc,
                title=intent.topic,
                subtitle=f"（{intent.doc_type}）" if intent.doc_type else "",
                date_str=datetime.datetime.now().strftime("%Y年%m月"),
            )

        # 写入内容
        self.markdown_to_docx(doc, content)

        # 文件名
        if not file_name:
            file_name = sanitize_filename(intent.topic)

        path = generate_sequential_filename(output_dir, file_name)
        doc.save(path)
        return path

    def export_news(self, title: str, source: str, publish_date: str,
                    content: str, output_dir: str, author: str = "",
                    file_name: str = None) -> str:
        """
        导出新闻文章为 .docx
        :param title: 文章标题
        :param source: 来源
        :param publish_date: 发布日期
        :param content: 正文
        :param output_dir: 输出目录
        :param author: 作者
        :param file_name: 文件名
        :return: 保存路径
        """
        os.makedirs(output_dir, exist_ok=True)

        doc = self.create_document(title)
        self.add_news_article(doc, title, source, publish_date, content, author)

        if not file_name:
            file_name = sanitize_filename(title)

        path = generate_sequential_filename(output_dir, file_name)
        doc.save(path)
        return path
